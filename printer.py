from docx import Document
from docx.shared import Mm, Pt
from pathlib import Path
from PIL import Image
from io import BytesIO
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import json
import threading

class TokenPrinterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("TokenPrinter")
        self.root.geometry("700x350")
        
        # Save settings in user's home directory
        home_dir = Path.home()
        self.settings_file = Path("tokenprinter_settings.json")
        
        self.load_settings()
        
        self.input_folder = self.settings.get("default_input", "")
        self.output_folder = self.settings.get("default_output", "")
        self.processing = False
        
        # Create main frame
        main_frame = tk.Frame(root, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Menu bar (Mac only)
        menubar = tk.Menu(root)
        root.config(menu=menubar)
        
        app_menu = tk.Menu(menubar, name='apple', tearoff=0)
        menubar.add_cascade(menu=app_menu)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Preferences...", command=self.open_settings, accelerator="Cmd+,")
        file_menu.add_separator()
        
        root.bind_all("<Command-comma>", lambda e: self.open_settings())
        
        # Input folder
        tk.Label(main_frame, text="Input Folder:", anchor='w').grid(row=0, column=0, padx=5, pady=10, sticky="w")
        self.input_entry = tk.Entry(main_frame, width=40)
        self.input_entry.grid(row=0, column=1, padx=5, pady=10, sticky="ew")
        self.input_entry.insert(0, self.input_folder)
        tk.Button(main_frame, text="Browse...", command=self.browse_input, width=12).grid(row=0, column=2, padx=5, pady=10)
        
        # Output folder
        tk.Label(main_frame, text="Output Folder:", anchor='w').grid(row=1, column=0, padx=5, pady=10, sticky="w")
        self.output_entry = tk.Entry(main_frame, width=40)
        self.output_entry.grid(row=1, column=1, padx=5, pady=10, sticky="ew")
        self.output_entry.insert(0, self.output_folder)
        tk.Button(main_frame, text="Browse...", command=self.browse_output, width=12).grid(row=1, column=2, padx=5, pady=10)
        
        # Filename
        tk.Label(main_frame, text="Filename:", anchor='w').grid(row=2, column=0, padx=5, pady=10, sticky="w")
        filename_frame = tk.Frame(main_frame)
        filename_frame.grid(row=2, column=1, columnspan=2, padx=5, pady=10, sticky="ew")
        self.filename_entry = tk.Entry(filename_frame, width=35)
        self.filename_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Label(filename_frame, text=".docx").pack(side=tk.LEFT, padx=(5, 0))
        
        # Configure column weights
        main_frame.columnconfigure(1, weight=1)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=3, column=0, columnspan=3, pady=10, sticky="ew")
        self.progress.grid_remove()
        
        # Status label
        self.status_label = tk.Label(main_frame, text="", fg="gray")
        self.status_label.grid(row=4, column=0, columnspan=3, pady=5)
        
        # Convert button
        self.convert_button = tk.Button(
            main_frame,
            text="Convert",
            command=self.convert,
            bg="#34C759",
            fg="black",
            activebackground="#30B350",
            activeforeground="black",
            font=("Arial", 14, "bold"),
            width=20,
            height=2,
            relief=tk.RAISED,
            bd=2
        )
        self.convert_button.grid(row=5, column=0, columnspan=3, pady=20)
    
    def load_settings(self):
        """Load settings from JSON file"""
        try:
            if self.settings_file.exists():
                with open(self.settings_file, 'r') as f:
                    self.settings = json.load(f)
            else:
                self.settings = {
                    "default_input": "",
                    "default_output": ""
                }
        except Exception:
            self.settings = {
                "default_input": "",
                "default_output": ""
            }
    
    def save_settings(self):
        """Save settings to JSON file"""
        try:
            with open(self.settings_file, 'w') as f:
                json.dump(self.settings, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error", f"Could not save settings:\n{str(e)}")
    
    def open_settings(self):
        """Open settings window"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Preferences")
        settings_window.geometry("700x250")
        
        settings_frame = tk.Frame(settings_window, padx=20, pady=20)
        settings_frame.pack(fill=tk.BOTH, expand=True)
        
        # Default input folder
        tk.Label(settings_frame, text="Default Input Folder:", anchor='w').grid(row=0, column=0, padx=5, pady=10, sticky="w")
        default_input_entry = tk.Entry(settings_frame, width=40)
        default_input_entry.grid(row=0, column=1, padx=5, pady=10, sticky="ew")
        default_input_entry.insert(0, self.settings.get("default_input", ""))
        tk.Button(settings_frame, text="Browse...", 
                  command=lambda: self.browse_settings_folder(default_input_entry), width=12).grid(row=0, column=2, padx=5, pady=10)
        
        # Default output folder
        tk.Label(settings_frame, text="Default Output Folder:", anchor='w').grid(row=1, column=0, padx=5, pady=10, sticky="w")
        default_output_entry = tk.Entry(settings_frame, width=40)
        default_output_entry.grid(row=1, column=1, padx=5, pady=10, sticky="ew")
        default_output_entry.insert(0, self.settings.get("default_output", ""))
        tk.Button(settings_frame, text="Browse...", 
                  command=lambda: self.browse_settings_folder(default_output_entry), width=12).grid(row=1, column=2, padx=5, pady=10)
        
        settings_frame.columnconfigure(1, weight=1)
        
        # Save button
        def save_and_close():
            self.settings["default_input"] = default_input_entry.get()
            self.settings["default_output"] = default_output_entry.get()
            self.save_settings()
            
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, self.settings["default_input"])
            self.input_folder = self.settings["default_input"]
            
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, self.settings["default_output"])
            self.output_folder = self.settings["default_output"]
            
            messagebox.showinfo("Success", "Settings saved!")
            settings_window.destroy()
        
        save_button = tk.Button(
            settings_frame,
            text="Save",
            command=save_and_close,
            bg="#007AFF",
            fg="black",
            activebackground="#0051D5",
            activeforeground="black",
            font=("Arial", 12, "bold"),
            width=15,
            height=2,
            relief=tk.RAISED,
            bd=2
        )
        save_button.grid(row=2, column=0, columnspan=3, pady=30)
    
    def browse_settings_folder(self, entry_widget):
        """Browse folder for settings window"""
        folder = filedialog.askdirectory(title="Select folder")
        if folder:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, folder)
    
    def browse_input(self):
        folder = filedialog.askdirectory(title="Select folder with images")
        if folder:
            self.input_folder = folder
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, folder)
    
    def browse_output(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.output_folder = folder
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, folder)
    
    def convert(self):
        if self.processing:
            return
        
        self.input_folder = self.input_entry.get()
        self.output_folder = self.output_entry.get()
        
        if not self.input_folder:
            messagebox.showerror("Error", "Please select an input folder")
            return
        
        if not self.output_folder:
            messagebox.showerror("Error", "Please select an output folder")
            return
        
        filename = self.filename_entry.get().strip()
        if not filename:
            messagebox.showerror("Error", "Please enter a filename")
            return
        
        self.processing = True
        self.convert_button.config(state='disabled', text='Processing...')
        self.progress.grid()
        self.progress.start(10)
        self.status_label.config(text="Processing images...")
        
        thread = threading.Thread(target=self.process_images, args=(filename,))
        thread.daemon = True
        thread.start()
    
    def process_images(self, filename):
        try:
            tokens_mappe = Path(self.input_folder)
            image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']
            images = [f for f in tokens_mappe.iterdir()
                      if f.is_file() and f.suffix.lower() in image_extensions]
            
            if not images:
                self.root.after(0, lambda: self.show_error("No images found in selected folder"))
                return
            
            doc = Document()
            
            section = doc.sections[0]
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(10)
            section.left_margin = Mm(10)
            section.right_margin = Mm(10)
            
            images_per_row = 2
            
            for i in range(0, len(images), images_per_row):
                self.root.after(0, lambda i=i, total=len(images): 
                               self.status_label.config(text=f"Processing image {i+1} of {total}..."))
                
                row_images = images[i:i + images_per_row]
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1
                
                for image in row_images:
                    img = Image.open(image)
                    rotated_img = img.rotate(-90, expand=True)
                   
                    if rotated_img.mode == 'RGBA':
                        rgb_img = Image.new('RGB', rotated_img.size, (255, 255, 255))
                        rgb_img.paste(rotated_img, mask=rotated_img.split()[3])
                        rotated_img = rgb_img
                   
                    img_buffer = BytesIO()
                    rotated_img.save(img_buffer, format='JPEG')
                    img_buffer.seek(0)
                   
                    run = p.add_run()
                    run.add_picture(img_buffer, width=Mm(88))
            
            self.root.after(0, lambda: self.status_label.config(text="Saving document..."))
            output_path = Path(self.output_folder) / f"{filename}.docx"
            doc.save(str(output_path))
            
            self.root.after(0, lambda: self.show_success(len(images), output_path))
            
        except Exception as e:
            self.root.after(0, lambda: self.show_error(f"An error occurred:\n{str(e)}"))
        finally:
            self.root.after(0, self.reset_ui)
    
    def show_success(self, count, path):
        self.progress.stop()
        self.progress.grid_remove()
        self.status_label.config(text="")
        messagebox.showinfo("Success", f"Processed {count} images!\nSaved to: {path}")
    
    def show_error(self, message):
        self.progress.stop()
        self.progress.grid_remove()
        self.status_label.config(text="")
        messagebox.showerror("Error", message)
    
    def reset_ui(self):
        self.processing = False
        self.convert_button.config(state='normal', text='Convert')
        self.progress.stop()
        self.progress.grid_remove()
        self.status_label.config(text="")

if __name__ == "__main__":
    root = tk.Tk()
    app = TokenPrinterApp(root)
    root.mainloop()